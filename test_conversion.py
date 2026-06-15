import os
import sys
import docx
from app import convert_docx_to_pdf, convert_pdf_to_docx

def main():
    print("Starting conversion integration test...")
    
    # 1. Create a dummy Word file
    doc_path = "test_input.docx"
    pdf_path = "test_output.pdf"
    docx_out_path = "test_output.docx"

    # Cleanup any old test outputs
    for path in [doc_path, pdf_path, docx_out_path]:
        if os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass

    print(f"1. Creating sample Word document: {doc_path}")
    doc = docx.Document()
    doc.add_heading("DocConvert Integration Test Document", 0)
    doc.add_paragraph("This is a simple test document created dynamically to test Word-to-PDF and PDF-to-Word conversions.")
    doc.add_paragraph("If you see this text in the converted PDF and DOCX files, the test has succeeded.")
    doc.save(doc_path)
    
    if not os.path.exists(doc_path):
        print("FAIL: Could not create test_input.docx")
        sys.exit(1)
    print("SUCCESS: test_input.docx created successfully.")

    # 2. Test Word to PDF
    print(f"2. Testing Word to PDF conversion: {doc_path} -> {pdf_path}")
    try:
        convert_docx_to_pdf(doc_path, pdf_path)
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            print("SUCCESS: Word to PDF conversion succeeded.")
        else:
            print("FAIL: test_output.pdf does not exist or is empty.")
            sys.exit(1)
    except Exception as e:
        print(f"FAIL: Word to PDF conversion failed with error: {e}")
        sys.exit(1)

    # 3. Test PDF to Word
    print(f"3. Testing PDF to Word conversion: {pdf_path} -> {docx_out_path}")
    try:
        convert_pdf_to_docx(pdf_path, docx_out_path)
        if os.path.exists(docx_out_path) and os.path.getsize(docx_out_path) > 0:
            print("SUCCESS: PDF to Word conversion succeeded.")
        else:
            print("FAIL: test_output.docx does not exist or is empty.")
            sys.exit(1)
    except Exception as e:
        print(f"FAIL: PDF to Word conversion failed with error: {e}")
        sys.exit(1)

    print("\n--- ALL CONVERSION TESTS PASSED SUCCESSFULLY! ---")
    
    # Cleanup files
    print("Cleaning up test files...")
    for path in [doc_path, pdf_path, docx_out_path]:
        if os.path.exists(path):
            os.remove(path)
    print("Cleanup done.")

if __name__ == "__main__":
    main()
